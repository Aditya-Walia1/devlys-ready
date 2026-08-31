from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Devlys_Smart_Review_QR_Operating_Guide.docx"

BLUE = "315EFB"
HEADING_BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "111827"
MUTED = "687387"
LIGHT_BLUE = "EEF2FF"
LIGHT_GREEN = "E7F6EE"
LIGHT_GOLD = "FFF5DD"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "111827"


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_shading(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        borders.append(left)
        p_pr.append(borders)


def add_numbering_definition(document, numbered=False):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if numbered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if numbered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(document, text, num_id, bold_prefix=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        remainder = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(remainder)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_callout(document, label, text, fill=LIGHT_BLUE, accent=BLUE):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.25
    set_paragraph_shading(paragraph, fill, accent)
    lead = paragraph.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=DARK_BLUE)
    body = paragraph.add_run(text)
    set_run_font(body, color=BLACK)
    return paragraph


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(document, text, bold_prefix=None):
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        lead = paragraph.add_run(bold_prefix)
        set_run_font(lead, bold=True)
        remainder = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(remainder)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(document, headers, rows, widths, alignments=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=10, bold=True, color=DARK_BLUE)
        if alignments:
            p.alignment = alignments[index]
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_run_font(run, size=9.5)
            if alignments:
                p.alignment = alignments[index]
    set_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, HEADING_BLUE, 18, 10),
        2: (13, HEADING_BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header_p.add_run("DEVLYS SMART REVIEW QR   |   OPERATING GUIDE")
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)
    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)


def build_document():
    document = Document()
    configure_document(document)
    bullets = add_numbering_definition(document, numbered=False)
    numbers = add_numbering_definition(document, numbered=True)

    # Editorial cover pattern with the compact_reference_guide preset.
    for _ in range(4):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run_font(kicker.add_run("OWNER & CUSTOMER HANDBOOK"), size=10, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("Devlys Smart Review QR"), size=30, color=NAVY, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_run_font(subtitle.add_run("How Devlys owners operate the service, how business customers use it, and how one QR stays reusable across six-month terms"), size=14, color=DARK_BLUE)

    audience = document.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    audience.paragraph_format.space_after = Pt(8)
    set_run_font(audience.add_run("Prepared for Devlys owners and approved Smart Review business customers"), size=10.5, color=MUTED, italic=True)
    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(version.add_run("Version 1.0  |  31 August 2026"), size=10, color=MUTED, bold=True)

    cover_note = document.add_paragraph()
    cover_note.paragraph_format.space_before = Pt(54)
    cover_note.paragraph_format.space_after = Pt(0)
    cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(cover_note.add_run("The customer writes. Devlys operates. The business benefits."), size=12, color=BLUE, bold=True)
    document.add_page_break()

    add_heading(document, "Purpose of this guide", 1)
    add_body(document, "This handbook explains the complete commercial operating model for Devlys Smart Review QR. It separates each role, shows the exact handoffs from enrollment through renewal, and provides practical checklists for running the service without confusing permissions or responsibilities.")
    add_callout(document, "Core rule", "Only approved Devlys owner accounts create businesses, locations and QR journeys. Business customers can view and use the QR codes assigned to them, submit payment references and monitor their own results. Reviewing customers never enter the management dashboard.")

    add_heading(document, "The three roles", 1)
    add_table(
        document,
        ["Role", "What this person does", "What this person cannot do"],
        [
            ("Devlys owner", "Approves applications, sets quotes and payment links, creates every QR location, verifies payment, activates service and renews terms.", "Cannot write or post a review on behalf of a reviewing customer."),
            ("Business customer", "Applies, pays Devlys, submits the payment reference, downloads and displays assigned QR codes, watches analytics and requests renewal.", "Cannot create a business, generate a new QR location, activate service or verify payment."),
            ("Reviewing customer", "Scans the QR, selects a rating and topics, adds a specific visit moment, edits the draft and decides whether to post on Google.", "Cannot access business analytics, subscriptions, payments or administration."),
        ],
        [1800, 4050, 3510],
    )

    add_heading(document, "Commercial lifecycle at a glance", 1)
    lifecycle = [
        "Business customer submits the public enrollment form with accurate business, contact, first-location and Google review details.",
        "Devlys owner reviews the application, confirms the plan, enters the six-month quote and adds an HTTPS payment link when available.",
        "Approval creates the client membership, first location and stable QR destination. The customer-facing route remains inactive until payment is verified.",
        "Business customer signs in using the same approved email, pays through the Devlys-provided route and submits the transaction reference.",
        "Devlys owner checks the real payment record and activates the business. The service start and end dates are set for six months.",
        "The business downloads and prints the Devlys-created QR. Customers scan it and choose whether to post an editable draft to Google.",
        "At renewal, the business pays again and submits a new reference. Devlys verifies it and extends the existing term by six months. The QR URL does not change.",
    ]
    for item in lifecycle:
        add_list_item(document, item, numbers)

    add_heading(document, "Part A - Devlys owner operating procedure", 1)
    add_callout(document, "Owner access", "Use an approved Devlys owner email. The initial owner allowlist includes the Devlys owner account configured for this Site. A signed-in user who is not an owner cannot approve applications or create QR locations.", fill=LIGHT_GREEN, accent="12764A")

    add_heading(document, "Open the owner dashboard", 2)
    for item in [
        "Open the Devlys website and select Dashboard.",
        "Sign in with the approved Devlys owner account.",
        "Confirm that the page labels you as Devlys owner and shows Commercial operations.",
        "Review the four overview totals: new applications, active businesses, payment checks and managed locations.",
    ]:
        add_list_item(document, item, numbers)

    add_heading(document, "Review a new enrollment", 2)
    add_body(document, "Every submitted business appears in the Enrollment queue. Before approval, Devlys owns the verification work.")
    for item in [
        "Confirm the business and contact names are credible and complete.",
        "Confirm the email is the exact address the business owner will use to sign in.",
        "Verify the phone or WhatsApp number through your normal sales process.",
        "Open the Google review link and confirm it belongs to the stated location.",
        "Confirm that the selected plan matches the required number of locations.",
    ]:
        add_list_item(document, item, bullets)
    add_callout(document, "Do not approve", "an application with a mismatched Google Business Profile, an email already assigned to another client, unclear ownership, or a payment link that is not HTTPS.", fill=LIGHT_GOLD, accent="C18700")

    add_heading(document, "Approve, quote and create the first QR", 2)
    for item in [
        "Enter the agreed total price in INR for the six-month term. The website stores the amount as the client quote; Devlys decides pricing outside the system.",
        "Paste the secure payment link supplied by the Devlys payment provider. This is optional in the software, but recommended for a clean client experience.",
        "Choose the QR brand colour and select Approve & create account.",
        "The system creates the business account, connects the approved email, creates the first location and generates one stable QR URL.",
        "The business status remains Pending payment. Scans do not open an active review journey until payment is verified.",
    ]:
        add_list_item(document, item, numbers)

    add_heading(document, "Verify payment and activate six months", 2)
    add_body(document, "The website records a payment reference; it does not independently prove that money reached the Devlys account. The Devlys owner must verify the payment in the real payment provider or bank before activation.")
    for item in [
        "Find the submitted payment under the correct business.",
        "Match the amount, payer, transaction reference and payment date against the Devlys payment account.",
        "Select Verify & add 6 months only after the payment is confirmed.",
        "Confirm the business status changes to Active and a Valid until date appears.",
        "Open the customer page once and confirm the location name and Google handoff are correct before delivering the QR artwork.",
    ]:
        add_list_item(document, item, numbers)

    add_heading(document, "Create additional QR locations", 2)
    add_body(document, "Only Devlys owners can create additional locations. Open the client account, expand Create another QR location, enter the location name, address, official Google review link and QR colour, then create the location.")
    add_table(
        document,
        ["Plan", "Maximum active locations", "Owner action when limit is reached"],
        [
            ("Starter", "1", "Upgrade the commercial plan before creating another location."),
            ("Growth", "3", "Confirm the new branch belongs to the same business account."),
            ("Scale", "10", "Review whether account-level support or custom terms are needed."),
        ],
        [1900, 2200, 5260],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    add_heading(document, "Deliver and support the QR", 2)
    for item in [
        "Use Download QR to obtain the 500 x 500 PNG.",
        "Send the PNG and the customer-page URL to the approved business contact.",
        "Advise the client to test one printed sample at the intended size before bulk printing.",
        "Explain that renewal keeps the same QR destination; do not create a new QR for a normal renewal.",
        "Monitor scans, drafts, Google handoffs and handoff rate for support conversations.",
    ]:
        add_list_item(document, item, bullets)

    add_heading(document, "Renew an existing business", 2)
    for item in [
        "Ask the business to use the current payment link or another approved Devlys payment route.",
        "The business submits the new transaction reference in its dashboard.",
        "Verify the payment externally, then select Verify & add 6 months.",
        "If the current term is still active, six months are added to the existing end date. If it has expired, the new term starts from the activation time.",
        "Do not replace the location slug or reprint the QR; the existing QR becomes active again after renewal.",
    ]:
        add_list_item(document, item, numbers)

    add_heading(document, "Part B - Business customer guide", 1)
    add_callout(document, "Customer responsibility", "Provide accurate information, pay Devlys through the approved route, submit the transaction reference, display only QR codes created by Devlys, monitor your results and renew on time. You do not create or activate QR codes yourself.", fill=LIGHT_BLUE, accent=BLUE)

    add_heading(document, "Enroll your business", 2)
    for item in [
        "Open Enroll a business from the Devlys website.",
        "Enter the legal or trading business name and the responsible contact person.",
        "Use the email address you will later use for dashboard sign-in.",
        "Enter the first location and paste its official Google review link from Google Business Profile.",
        "Choose the plan that matches the expected number of locations and submit the form.",
        "Save the application reference shown on the confirmation screen.",
    ]:
        add_list_item(document, item, numbers)

    add_heading(document, "Sign in and pay", 2)
    for item in [
        "Wait for Devlys to approve the application and confirm the quote.",
        "Open Dashboard and sign in using exactly the approved email address.",
        "Review the plan, six-month quote, payment status and payment link.",
        "Pay through the Devlys-provided secure route. Never pay to account details received from an unverified sender.",
        "Return to the dashboard, select the payment method and enter the UTR, UPI reference or receipt ID.",
        "Select Submit for verification. The service activates only after a Devlys owner confirms receipt.",
    ]:
        add_list_item(document, item, numbers)

    add_heading(document, "Download, print and place the QR", 2)
    add_body(document, "After activation, each assigned location appears under Your locations. Use Download QR for print artwork or Copy URL for digital sharing.")
    for item in [
        "Keep enough white space around the QR and do not crop or stretch it.",
        "Print a test copy and scan it from the expected customer distance.",
        "Place it where a real customer can respond after the experience: table card, counter, receipt, bill folder or follow-up message.",
        "Do not offer rewards only for positive reviews and do not ask staff to write reviews for customers.",
        "If a location name or Google review link is wrong, contact Devlys instead of printing a replacement QR yourself.",
    ]:
        add_list_item(document, item, bullets)

    add_heading(document, "Read the analytics", 2)
    add_table(
        document,
        ["Metric", "Meaning", "Useful interpretation"],
        [
            ("QR scans", "Number of times the customer journey was opened in the last 30 days.", "Low scans usually point to QR placement, visibility or staff adoption."),
            ("Drafts created", "Customers who supplied a rating and specific visit moment, then generated an editable draft.", "Compare with scans to understand whether the form is easy and relevant."),
            ("Google handoffs", "Customers who selected Copy and continue to Google.", "This is intent to visit Google, not proof that a review was posted."),
            ("Handoff rate", "Google handoffs divided by QR scans.", "Use as a funnel signal; do not treat it as a published-review rate."),
        ],
        [1600, 3680, 4080],
    )

    add_heading(document, "Renew without replacing the QR", 2)
    add_body(document, "Before the Valid until date, pay the next term and submit the new transaction reference. Devlys verifies the payment and adds six months. Every existing printed QR continues to point to the same URL. If the term expires first, the customer page pauses until Devlys completes renewal.")

    document.add_page_break()
    add_heading(document, "Part C - What the reviewing customer experiences", 1)
    for item in [
        "Scan the business location QR. No management sign-in is required.",
        "Choose a star rating that honestly reflects the visit.",
        "Select relevant topics such as food, service, ambience or value.",
        "Write at least one specific moment from the visit. A concrete detail produces a more natural draft than a generic compliment.",
        "Read the generated paragraph and edit it until it sounds accurate and personal.",
        "Select Copy and continue to Google. The draft is copied and the official Google review page opens.",
        "Choose whether to paste, edit and post. Devlys never posts automatically.",
    ]:
        add_list_item(document, item, numbers)
    add_callout(document, "AI writing standard", "The service uses the customer’s rating, selected topics and specific moment. The prompt asks for natural language, avoids generic marketing clichés, preserves criticism for low ratings and prohibits invented details. If the AI service is unavailable, a safe fallback still uses the supplied moment and rating.", fill=LIGHT_GREEN, accent="12764A")

    add_heading(document, "Status glossary", 1)
    add_table(
        document,
        ["Status", "Meaning and next action"],
        [
            ("Application submitted", "Devlys must review the business and Google details."),
            ("Application approved", "The account and first stable QR destination exist; payment or activation may still be pending."),
            ("Payment unpaid", "The client has not submitted a payment reference."),
            ("Payment submitted", "The client entered a reference; Devlys must verify it externally."),
            ("Payment paid", "A Devlys owner verified the payment."),
            ("Business pending payment", "The QR destination exists but the customer review journey is not active."),
            ("Business active", "The customer QR journey works until the service end date."),
            ("Business expired", "The end date has passed. Renewal reactivates the same QR destination."),
        ],
        [2200, 7160],
    )

    add_heading(document, "Troubleshooting", 1)
    troubleshooting = [
        ("I signed in but see No linked business", "Use the exact email entered during enrollment. If it is correct, ask a Devlys owner to check the client membership."),
        ("The QR says the service is inactive", "The business is awaiting payment verification or the six-month term has expired. Devlys must activate or renew it."),
        ("The QR opens the wrong Google listing", "Stop distributing the QR and ask Devlys to correct the location’s official Google review link."),
        ("A payment reference is submitted but service is pending", "Devlys has not yet verified the transaction against the real payment account."),
        ("The draft sounds too generic", "Add a more concrete moment: name the interaction, item, timing, outcome or detail that actually occurred. Do not invent information."),
        ("A client needs another QR", "Only a Devlys owner can add it, and the selected plan must have an available location slot."),
    ]
    for problem, action in troubleshooting:
        add_body(document, f"{problem}: {action}", bold_prefix=f"{problem}: ")

    add_heading(document, "Launch prerequisites", 1)
    add_callout(document, "Current hosting note", "An owner-only private Site is suitable for internal review but prevents real customers from opening public QR routes. Before commercial launch, publish the Site for public visitors while keeping dashboard and management APIs protected by the server-side Devlys owner allowlist and client membership checks.", fill=LIGHT_GOLD, accent="C18700")
    for item in [
        "Confirm the public enrollment page and every active QR route can be opened without management access.",
        "Confirm only approved owner emails can open commercial operations or call admin APIs.",
        "Configure the OpenAI API key in Sites if AI-generated drafts are required; otherwise the safe fallback is used.",
        "Choose and configure the real Devlys payment provider and provide an HTTPS payment link for each approved client.",
        "Complete one live payment verification and one live QR-to-Google test before selling the service.",
        "Set the custom .com domain, privacy notice, terms, billing contact and support contact before public launch.",
    ]:
        add_list_item(document, item, bullets)

    add_heading(document, "Owner go-live checklist", 1)
    for item in [
        "Application details and Google listing verified",
        "Correct plan and six-month quote confirmed",
        "Secure payment route sent to the client",
        "Payment received and externally verified",
        "Business status Active with correct end date",
        "QR downloaded and scanned from a printed test",
        "Customer page content and Google handoff checked",
        "Client understands analytics, renewal and support route",
    ]:
        add_list_item(document, item, bullets)

    add_heading(document, "Business customer checklist", 1)
    for item in [
        "Enrollment submitted with the email used for dashboard sign-in",
        "Quote and payment destination checked",
        "Transaction reference submitted",
        "Activation confirmed before printing",
        "QR test-scanned at the intended print size",
        "QR displayed only at the matching business location",
        "Analytics reviewed regularly",
        "Renewal started before the Valid until date",
    ]:
        add_list_item(document, item, bullets)

    add_callout(document, "Final operating principle", "Devlys controls provisioning and activation. The business controls placement and renewal. The reviewing customer controls the words and the final decision to post.", fill=LIGHT_BLUE, accent=BLUE)

    document.core_properties.title = "Devlys Smart Review QR Operating Guide"
    document.core_properties.subject = "Owner and business customer operating handbook"
    document.core_properties.author = "Devlys"
    document.core_properties.keywords = "Devlys, Smart Review QR, owner guide, customer guide, six-month service"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
